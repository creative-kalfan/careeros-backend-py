"""Native OOXML DOCX Compiler for CareerOS Resume Studio.

Compiles the canonical ResumeDocumentModel and DocumentStyleModel into a genuine,
fully editable Microsoft Word document (.docx) using native OpenXML constructs.
"""

from __future__ import annotations

import io
from typing import Optional

from docx import Document  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.shared import Inches, Pt, RGBColor  # type: ignore

from .document_model import (
    BulletElement,
    DocumentElement,
    EducationEntry,
    ExperiencePosition,
    ProjectEntry,
    ResumeDocumentModel,
    SkillGroup,
)
from .style_model import DocumentStyleModel, hex_to_rgb


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set inner margins for a table cell in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_paragraph_bottom_border(paragraph, color_hex: str = "CBD5E1", size: int = 6):
    """Add a native OOXML bottom border to a paragraph (section divider line)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


class DocxCompiler:
    """Compiles ResumeDocumentModel into a native editable OOXML DOCX."""

    def compile(self, doc_model: ResumeDocumentModel) -> bytes:
        style = doc_model.style or DocumentStyleModel()
        doc = Document()

        # 1. Page settings & margins
        section = doc.sections[0]
        section.page_width = Pt(style.page_width_pt)
        section.page_height = Pt(style.page_height_pt)
        section.top_margin = Pt(style.margin_top_pt)
        section.bottom_margin = Pt(style.margin_bottom_pt)
        section.left_margin = Pt(style.margin_left_pt)
        section.right_margin = Pt(style.margin_right_pt)

        # Base typography colors
        body_rgb = RGBColor(*hex_to_rgb(style.body_color_hex))
        heading_rgb = RGBColor(*hex_to_rgb(style.heading_color_hex))
        accent_rgb = RGBColor(*hex_to_rgb(style.accent_color_hex))
        meta_rgb = RGBColor(*hex_to_rgb(style.meta_color_hex))

        # 2. Header (Name, Headline, Contact details)
        hdr = doc_model.header
        if hdr.full_name:
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_before = Pt(0)
            p_name.paragraph_format.space_after = Pt(1.5)
            r_name = p_name.add_run(hdr.full_name)
            r_name.font.name = style.heading_font
            r_name.font.size = Pt(style.name_size_pt)
            r_name.font.bold = True
            r_name.font.color.rgb = heading_rgb

        if hdr.headline:
            p_head = doc.add_paragraph()
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_head.paragraph_format.space_before = Pt(0)
            p_head.paragraph_format.space_after = Pt(2.0)
            r_head = p_head.add_run(hdr.headline)
            r_head.font.name = style.body_font
            r_head.font.size = Pt(style.headline_size_pt)
            r_head.font.bold = True
            r_head.font.color.rgb = accent_rgb

        contact_text = hdr.contact_line()
        if contact_text:
            p_contact = doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.paragraph_format.space_before = Pt(0)
            p_contact.paragraph_format.space_after = Pt(style.section_after_pt)
            r_contact = p_contact.add_run(contact_text)
            r_contact.font.name = style.body_font
            r_contact.font.size = Pt(style.meta_size_pt)
            r_contact.font.color.rgb = meta_rgb

        # Helper for adding standard section headings
        def add_section_heading(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(style.section_before_pt)
            p.paragraph_format.space_after = Pt(style.section_after_pt)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title.upper())
            run.font.name = style.heading_font
            run.font.size = Pt(style.section_heading_size_pt)
            run.font.bold = True
            run.font.color.rgb = heading_rgb
            if style.heading_has_divider:
                _set_paragraph_bottom_border(p, color_hex=style.divider_color_hex)
            return p

        # 3. Dynamic Section Rendering
        for sec in doc_model.section_order:
            if sec == "summary" and doc_model.summary and doc_model.summary.text.strip():
                add_section_heading("Professional Summary")
                p_sum = doc.add_paragraph()
                p_sum.paragraph_format.space_before = Pt(2.0)
                p_sum.paragraph_format.space_after = Pt(style.paragraph_after_pt)
                p_sum.paragraph_format.line_spacing = style.line_spacing
                r_sum = p_sum.add_run(doc_model.summary.text.strip())
                r_sum.font.name = style.body_font
                r_sum.font.size = Pt(style.body_size_pt)
                r_sum.font.color.rgb = body_rgb

            elif sec == "experience" and doc_model.experience:
                add_section_heading("Experience")
                for exp in doc_model.experience:
                    # Header row: Role at Company (Left), Date | Location (Right)
                    usable_width_in = (style.page_width_pt - style.margin_left_pt - style.margin_right_pt) / 72.0
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = False
                    col_widths = (Inches(usable_width_in * 0.65), Inches(usable_width_in * 0.35))

                    cell_left, cell_right = tbl.rows[0].cells
                    cell_left.width = col_widths[0]
                    cell_right.width = col_widths[1]
                    _set_cell_margins(cell_left, top=30, bottom=30, left=0, right=0)
                    _set_cell_margins(cell_right, top=30, bottom=30, left=0, right=0)

                    # Left: Role and Company
                    p_l = cell_left.paragraphs[0]
                    p_l.paragraph_format.space_before = Pt(0)
                    p_l.paragraph_format.space_after = Pt(1.0)
                    r_role = p_l.add_run(exp.role)
                    r_role.font.name = style.body_font
                    r_role.font.size = Pt(style.subheading_size_pt)
                    r_role.font.bold = True
                    r_role.font.color.rgb = heading_rgb

                    if exp.company:
                        r_at = p_l.add_run(f" | {exp.company}")
                        r_at.font.name = style.body_font
                        r_at.font.size = Pt(style.subheading_size_pt)
                        r_at.font.bold = False
                        r_at.font.color.rgb = body_rgb

                    # Right: Dates and Location
                    p_r = cell_right.paragraphs[0]
                    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_r.paragraph_format.space_before = Pt(0)
                    p_r.paragraph_format.space_after = Pt(1.0)
                    meta_parts = [exp.date_range, exp.location]
                    meta_str = " | ".join(filter(None, meta_parts))
                    r_meta = p_r.add_run(meta_str)
                    r_meta.font.name = style.body_font
                    r_meta.font.size = Pt(style.meta_size_pt)
                    r_meta.font.color.rgb = meta_rgb

                    # Bullets
                    for b in exp.bullets:
                        if not b.text.strip():
                            continue
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_before = Pt(0)
                        p_b.paragraph_format.space_after = Pt(1.5)
                        p_b.paragraph_format.line_spacing = style.line_spacing
                        r_b = p_b.add_run(b.text.strip())
                        r_b.font.name = style.body_font
                        r_b.font.size = Pt(style.body_size_pt)
                        r_b.font.color.rgb = body_rgb

            elif sec == "internships" and doc_model.internships:
                add_section_heading("Internships")
                for exp in doc_model.internships:
                    usable_width_in = (style.page_width_pt - style.margin_left_pt - style.margin_right_pt) / 72.0
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = False
                    col_widths = (Inches(usable_width_in * 0.65), Inches(usable_width_in * 0.35))
                    cell_left, cell_right = tbl.rows[0].cells
                    cell_left.width = col_widths[0]
                    cell_right.width = col_widths[1]

                    p_l = cell_left.paragraphs[0]
                    r_role = p_l.add_run(exp.role)
                    r_role.font.bold = True
                    if exp.company:
                        p_l.add_run(f" | {exp.company}")

                    p_r = cell_right.paragraphs[0]
                    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_r.add_run(exp.date_range)

                    for b in exp.bullets:
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(1.5)
                        r_b = p_b.add_run(b.text.strip())
                        r_b.font.size = Pt(style.body_size_pt)

            elif sec == "projects" and doc_model.projects:
                add_section_heading("Projects")
                for prj in doc_model.projects:
                    p_prj = doc.add_paragraph()
                    p_prj.paragraph_format.space_before = Pt(2.0)
                    p_prj.paragraph_format.space_after = Pt(1.0)
                    r_pname = p_prj.add_run(prj.name)
                    r_pname.font.name = style.body_font
                    r_pname.font.size = Pt(style.subheading_size_pt)
                    r_pname.font.bold = True
                    r_pname.font.color.rgb = heading_rgb

                    if prj.technologies:
                        r_tech = p_prj.add_run(f" ({', '.join(prj.technologies)})")
                        r_tech.font.name = style.body_font
                        r_tech.font.size = Pt(style.meta_size_pt)
                        r_tech.font.italic = True
                        r_tech.font.color.rgb = meta_rgb

                    if prj.description:
                        p_desc = doc.add_paragraph(style='List Bullet')
                        p_desc.paragraph_format.space_before = Pt(0)
                        p_desc.paragraph_format.space_after = Pt(1.5)
                        p_desc.paragraph_format.line_spacing = style.line_spacing
                        r_desc = p_desc.add_run(prj.description)
                        r_desc.font.name = style.body_font
                        r_desc.font.size = Pt(style.body_size_pt)
                        r_desc.font.color.rgb = body_rgb

                    for b in prj.bullets:
                        if not b.text.strip() or b.text.strip() == prj.description.strip():
                            continue
                        p_b = doc.add_paragraph(style='List Bullet')
                        p_b.paragraph_format.space_after = Pt(1.5)
                        r_b = p_b.add_run(b.text.strip())
                        r_b.font.name = style.body_font
                        r_b.font.size = Pt(style.body_size_pt)
                        r_b.font.color.rgb = body_rgb

            elif sec == "education" and doc_model.education:
                add_section_heading("Education")
                for edu in doc_model.education:
                    usable_width_in = (style.page_width_pt - style.margin_left_pt - style.margin_right_pt) / 72.0
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.autofit = False
                    col_widths = (Inches(usable_width_in * 0.7), Inches(usable_width_in * 0.3))
                    cell_left, cell_right = tbl.rows[0].cells
                    cell_left.width = col_widths[0]
                    cell_right.width = col_widths[1]
                    p_l = cell_left.paragraphs[0]

                    f_study = getattr(edu, "field_of_study", getattr(edu, "field", ""))
                    degree_str = " in ".join(filter(None, [edu.degree, f_study])) or "Degree"
                    r_deg = p_l.add_run(degree_str)
                    r_deg.font.name = style.body_font
                    r_deg.font.size = Pt(style.subheading_size_pt)
                    r_deg.font.bold = True
                    r_deg.font.color.rgb = heading_rgb

                    if edu.institution:
                        r_inst = p_l.add_run(f" — {edu.institution}")
                        r_inst.font.name = style.body_font
                        r_inst.font.size = Pt(style.subheading_size_pt)
                        r_inst.font.color.rgb = body_rgb

                    p_r = cell_right.paragraphs[0]
                    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r_date = p_r.add_run(edu.date_range)
                    r_date.font.name = style.body_font
                    r_date.font.size = Pt(style.meta_size_pt)
                    r_date.font.color.rgb = meta_rgb

                    if edu.coursework:
                        p_cw = doc.add_paragraph()
                        p_cw.paragraph_format.space_before = Pt(0)
                        p_cw.paragraph_format.space_after = Pt(2.0)
                        r_cw = p_cw.add_run(f"Coursework: {', '.join(edu.coursework)}")
                        r_cw.font.name = style.body_font
                        r_cw.font.size = Pt(style.meta_size_pt)
                        r_cw.font.color.rgb = meta_rgb

            elif sec == "skills" and doc_model.skills:
                add_section_heading("Skills")
                for grp in doc_model.skills:
                    if not grp.skills:
                        continue
                    p_sk = doc.add_paragraph()
                    p_sk.paragraph_format.space_before = Pt(1.0)
                    p_sk.paragraph_format.space_after = Pt(2.0)
                    p_sk.paragraph_format.line_spacing = style.line_spacing
                    r_cat = p_sk.add_run(f"{grp.category}: ")
                    r_cat.font.name = style.body_font
                    r_cat.font.size = Pt(style.body_size_pt)
                    r_cat.font.bold = True
                    r_cat.font.color.rgb = heading_rgb

                    r_vals = p_sk.add_run(", ".join(grp.skills))
                    r_vals.font.name = style.body_font
                    r_vals.font.size = Pt(style.body_size_pt)
                    r_vals.font.color.rgb = body_rgb

            elif sec == "certifications" and doc_model.certifications:
                add_section_heading("Certifications")
                for c in doc_model.certifications:
                    p_c = doc.add_paragraph(style='List Bullet')
                    p_c.paragraph_format.space_after = Pt(1.5)
                    r_c = p_c.add_run(c.text)
                    r_c.font.name = style.body_font
                    r_c.font.size = Pt(style.body_size_pt)
                    r_c.font.color.rgb = body_rgb

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


docx_compiler = DocxCompiler()
