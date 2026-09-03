"""DOCX parser using python-docx with style-aware extraction."""

from __future__ import annotations

import logging
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .models import DocumentBlock, DocumentLine, DocumentSpan, ParsedResume, ParseResult
from .section_detector import detect_sections, get_section_blocks
from .contact_parser import extract_contact_from_blocks
from .experience_parser import parse_experience_section
from .education_parser import parse_education_section
from .skills_parser import parse_skills_section
from .project_parser import parse_projects_section
from .other_parsers import (
    parse_certifications,
    parse_achievements,
    parse_languages,
    parse_links,
    parse_summary,
)

logger = logging.getLogger(__name__)


class DOCXParser:
    """DOCX parser with style-aware extraction."""

    def __init__(self, debug: bool = False):
        self.debug = debug
        self.parse_notes: List[str] = []

    def parse(self, file_path: str) -> ParseResult:
        """Parse a DOCX file and return structured resume data."""
        self.parse_notes = []
        
        try:
            doc = Document(file_path)
            logger.info("DOCX opened: %d paragraphs", len(doc.paragraphs))

            all_blocks: List[DocumentBlock] = []
            raw_text_parts: List[str] = []

            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)

                # Extract runs for font info
                spans = self._extract_spans_from_paragraph(para, para_idx)
                
                if spans:
                    # Each paragraph becomes a line
                    line = DocumentLine(
                        spans=spans,
                        page=0,  # DOCX doesn't have pages
                        x0=min(s.x0 for s in spans),
                        y0=para_idx * 20,  # Approximate y position
                        x1=max(s.x1 for s in spans),
                        y1=para_idx * 20 + max(s.font_size for s in spans),
                    )
                    # Each paragraph becomes a block (single line)
                    block = DocumentBlock(
                        lines=[line],
                        page=0,
                        x0=line.x0, y0=line.y0, x1=line.x1, y1=line.y1,
                    )
                    all_blocks.append(block)

            raw_text = "\n".join(raw_text_parts)

            if not all_blocks:
                return ParseResult(
                    status="failed",
                    error="No text content extracted from DOCX",
                    raw_text=raw_text,
                    debug_info={"parse_notes": self.parse_notes} if self.debug else None,
                )

            # For DOCX, we don't have multi-page layout, create a simple page layout
            from .layout import PageLayout
            page_layout = PageLayout(
                page_num=0,
                page_width=612,  # Standard letter width in points
                page_height=792,
                columns=[],
                is_multi_column=False,
                body_font_size=11.0,
            )
            
            # Assign all blocks to page 0
            for block in all_blocks:
                block.page = 0
                for line in block.lines:
                    line.page = 0
                    for span in line.spans:
                        span.page = 0

            page_layouts = [page_layout]

            # Parse structured content
            parsed = self._parse_structured(all_blocks, page_layouts, raw_text)

            return ParseResult(
                status="completed",
                parsed=parsed,
                raw_text=raw_text,
                debug_info={"parse_notes": self.parse_notes} if self.debug else None,
            )

        except Exception as e:
            logger.exception("DOCX parsing failed")
            return ParseResult(
                status="failed",
                error=f"DOCX parsing failed: {str(e)}",
                raw_text="",
                debug_info={"parse_notes": self.parse_notes} if self.debug else None,
            )

    def _extract_spans_from_paragraph(self, para: Paragraph, para_idx: int) -> List[DocumentSpan]:
        """Extract spans from a DOCX paragraph."""
        spans = []
        
        for run_idx, run in enumerate(para.runs):
            text = run.text
            if not text:
                continue

            # Get font properties
            font_size = 11.0
            if run.font.size:
                font_size = run.font.size.pt

            bold = run.bold if run.bold is not None else False
            italic = run.italic if run.italic is not None else False
            font_name = run.font.name if run.font.name else ""

            # Approximate positions
            x0 = run_idx * 100
            y0 = para_idx * 20
            x1 = x0 + len(text) * font_size * 0.5
            y1 = y0 + font_size

            spans.append(DocumentSpan(
                text=text,
                page=0,
                x0=x0,
                y0=y0,
                x1=x1,
                y1=y1,
                font_size=font_size,
                bold=bold,
                italic=italic,
                font_name=font_name,
                color=0,
            ))

        return spans

    def _parse_structured(
        self,
        blocks: List[DocumentBlock],
        page_layouts: List[PageLayout],
        raw_text: str,
    ) -> ParsedResume:
        """Parse structured content from blocks."""
        # Detect sections
        sections, section_notes = detect_sections(blocks, page_layouts)
        self.parse_notes.extend(section_notes)

        # Extract contact
        contact = extract_contact_from_blocks(blocks)

        # Parse each section
        experience_result = parse_experience_section(get_section_blocks(sections, "experience"))
        education_result = parse_education_section(get_section_blocks(sections, "education"))
        skills = parse_skills_section(get_section_blocks(sections, "skills"))
        projects_result = parse_projects_section(get_section_blocks(sections, "projects"))
        certifications = parse_certifications(get_section_blocks(sections, "certifications"))
        achievements = parse_achievements(get_section_blocks(sections, "achievements"))
        languages = parse_languages(get_section_blocks(sections, "languages"))
        links = parse_links(get_section_blocks(sections, "links"))
        summary = parse_summary(get_section_blocks(sections, "summary"))

        self.parse_notes.extend(experience_result.parse_notes)
        self.parse_notes.extend(education_result.parse_notes)
        self.parse_notes.extend(projects_result.parse_notes)

        return ParsedResume(
            contact=contact,
            summary=summary if summary else None,
            experience=experience_result.experience,
            education=education_result.education,
            skills=skills,
            projects=projects_result.projects,
            certifications=certifications,
            achievements=achievements,
            languages=languages,
            links=links,
            parse_notes=self.parse_notes,
        )