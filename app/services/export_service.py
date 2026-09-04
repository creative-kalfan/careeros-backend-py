"""Resume export service for PDF and DOCX generation (Step 6)."""

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Any

from app.models.resume import ResumeContent

logger = logging.getLogger(__name__)


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_\- ]+", "", name)
    cleaned = re.sub(r"\s+", "_", cleaned.strip())
    return cleaned[:80] or "resume"


def _render_html(content: ResumeContent, template: str = "minimal") -> str:
    profile = content.profile
    personal = profile.personal
    sections: list[str] = []

    # Header
    name = personal.full_name or "Candidate"
    headline = personal.headline or profile.target_role or ""
    header_html = f'<div class="header"><h1>{name}</h1>'
    if headline:
        header_html += f'<p class="headline">{headline}</p>'
    header_html += '</div>'
    sections.append(header_html)

    # Contact info bar
    contact_items = []
    if personal.email:
        contact_items.append(personal.email)
    if personal.phone:
        contact_items.append(personal.phone)
    if personal.location:
        contact_items.append(personal.location)
    if personal.linkedin:
        contact_items.append(personal.linkedin)
    if personal.github:
        contact_items.append(personal.github)
    if personal.website:
        contact_items.append(personal.website)

    if contact_items:
        sections.append(f'<div class="contact">{" | ".join(contact_items)}</div>')

    # Professional Summary
    if profile.summary:
        sections.append(
            f'<div class="section">'
            f'<h2 class="section-title">Professional Summary</h2>'
            f'<p class="summary-text">{profile.summary}</p>'
            f'</div>'
        )

    # Experience
    if profile.experience:
        items = []
        for exp in profile.experience:
            role = exp.role or "Role"
            company = exp.company or ""
            title_str = f"<strong>{role}</strong>" + (f" at {company}" if company else "")
            date_range = " — ".join(filter(None, [exp.start_date, "Present" if exp.current else exp.end_date]))
            meta_parts = [date_range]
            if exp.location:
                meta_parts.append(exp.location)
            meta_str = " | ".join(filter(None, meta_parts))

            bullets_html = ""
            bullet_texts = exp.get_all_bullet_texts()
            if bullet_texts:
                lis = "".join(f"<li>{b}</li>" for b in bullet_texts if b.strip())
                if lis:
                    bullets_html = f'<ul class="bullet-list">{lis}</ul>'

            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{title_str}</span><span class="entry-meta">{meta_str}</span></div>'
                f'{bullets_html}'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Experience</h2>{"".join(items)}</div>')

    # Internships
    if profile.internships:
        items = []
        for exp in profile.internships:
            role = exp.role or "Intern"
            company = exp.company or ""
            title_str = f"<strong>{role}</strong>" + (f" at {company}" if company else "")
            date_range = " — ".join(filter(None, [exp.start_date, "Present" if exp.current else exp.end_date]))
            meta_parts = [date_range]
            if exp.location:
                meta_parts.append(exp.location)
            meta_str = " | ".join(filter(None, meta_parts))

            bullets_html = ""
            bullet_texts = exp.get_all_bullet_texts()
            if bullet_texts:
                lis = "".join(f"<li>{b}</li>" for b in bullet_texts if b.strip())
                if lis:
                    bullets_html = f'<ul class="bullet-list">{lis}</ul>'

            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{title_str}</span><span class="entry-meta">{meta_str}</span></div>'
                f'{bullets_html}'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Internships</h2>{"".join(items)}</div>')

    # Projects
    if profile.projects:
        items = []
        for proj in profile.projects:
            name_str = f"<strong>{proj.name or 'Project'}</strong>"
            if proj.url:
                name_str += f' <span class="entry-url">({proj.url})</span>'
            tech_str = f'<em>Technologies: {", ".join(proj.technologies)}</em>' if proj.technologies else ""

            bullets = []
            if proj.description:
                bullets.append(proj.description)
            if proj.problem:
                bullets.append(f"Problem: {proj.problem}")
            if proj.contribution:
                bullets.append(f"Contribution: {proj.contribution}")
            if proj.results:
                bullets.append(f"Results: {proj.results}")

            bullets_html = ""
            if bullets:
                lis = "".join(f"<li>{b}</li>" for b in bullets if b.strip())
                bullets_html = f'<ul class="bullet-list">{lis}</ul>'

            tech_html = f'<p class="project-tech">{tech_str}</p>' if tech_str else ""

            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{name_str}</span></div>'
                f'{bullets_html}'
                f'{tech_html}'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Projects</h2>{"".join(items)}</div>')

    # Education
    if profile.education:
        items = []
        for edu in profile.education:
            degree_str = " in ".join(filter(None, [edu.degree, edu.field])) or "Degree"
            inst_str = f"<strong>{degree_str}</strong>" + (f" — {edu.institution}" if edu.institution else "")
            date_range = " — ".join(filter(None, [edu.start_date, edu.end_date]))
            meta_parts = [date_range]
            if edu.location:
                meta_parts.append(edu.location)
            if edu.gpa:
                meta_parts.append(f"GPA: {edu.gpa}")
            meta_str = " | ".join(filter(None, meta_parts))

            extra_html = ""
            if edu.coursework:
                extra_html += f'<p class="edu-detail">Coursework: {", ".join(edu.coursework)}</p>'
            if edu.achievements:
                lis = "".join(f"<li>{a}</li>" for a in edu.achievements if a.strip())
                if lis:
                    extra_html += f'<ul class="bullet-list">{lis}</ul>'

            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{inst_str}</span><span class="entry-meta">{meta_str}</span></div>'
                f'{extra_html}'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Education</h2>{"".join(items)}</div>')

    # Skills
    if profile.skills and any([
        profile.skills.technical,
        profile.skills.tools,
        profile.skills.languages,
        profile.skills.databases,
        profile.skills.analytics,
        profile.skills.soft_skills,
        profile.skills.custom,
    ]):
        items = []
        categories = [
            ("technical", "Technical Skills"),
            ("tools", "Tools & Frameworks"),
            ("languages", "Languages"),
            ("databases", "Databases"),
            ("analytics", "Analytics"),
            ("soft_skills", "Soft Skills"),
        ]
        for cat_key, cat_label in categories:
            vals = getattr(profile.skills, cat_key, [])
            if vals:
                items.append(
                    f'<div class="skill-row"><strong>{cat_label}:</strong> {", ".join(vals)}</div>'
                )
        if profile.skills.custom:
            for custom_label, custom_vals in profile.skills.custom.items():
                if custom_vals:
                    items.append(
                        f'<div class="skill-row"><strong>{custom_label}:</strong> {", ".join(custom_vals)}</div>'
                    )
        if items:
            sections.append(f'<div class="section"><h2 class="section-title">Skills</h2>{"".join(items)}</div>')

    # Certifications
    if profile.certifications:
        items = []
        for cert in profile.certifications:
            name_str = f"<strong>{cert.name or 'Certification'}</strong>"
            details = " | ".join(filter(None, [cert.issuer, cert.date]))
            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{name_str}</span><span class="entry-meta">{details}</span></div>'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Certifications</h2>{"".join(items)}</div>')

    # Achievements
    if profile.achievements:
        lis = "".join(f"<li>{ach}</li>" for ach in profile.achievements if ach.strip())
        if lis:
            sections.append(f'<div class="section"><h2 class="section-title">Achievements</h2><ul class="bullet-list">{lis}</ul></div>')

    # Leadership
    if profile.leadership:
        items = []
        for lead in profile.leadership:
            header = f"<strong>{lead.role or 'Leader'}</strong>" + (f" — {lead.organization}" if lead.organization else "")
            date_range = " — ".join(filter(None, [lead.start_date, lead.end_date]))
            desc = f'<p class="lead-desc">{lead.description}</p>' if lead.description else ""
            items.append(
                f'<div class="entry">'
                f'<div class="entry-header"><span class="entry-title">{header}</span><span class="entry-meta">{date_range}</span></div>'
                f'{desc}'
                f'</div>'
            )
        sections.append(f'<div class="section"><h2 class="section-title">Leadership</h2>{"".join(items)}</div>')

    # Languages
    if profile.languages:
        items = []
        for lang in profile.languages:
            label = lang.language or ""
            if lang.proficiency:
                label += f" ({lang.proficiency})"
            if label:
                items.append(label)
        if items:
            sections.append(f'<div class="section"><h2 class="section-title">Languages</h2><p class="languages-list">{", ".join(items)}</p></div>')

    # Links
    if profile.links:
        items = []
        for link in profile.links:
            label = link.label or link.url or ""
            items.append(f'<div class="link-item"><strong>{label}:</strong> {link.url or ""}</div>')
        if items:
            sections.append(f'<div class="section"><h2 class="section-title">Links</h2>{"".join(items)}</div>')

    # Additional
    if profile.additional:
        items = []
        for add in profile.additional:
            if add.title:
                items.append(f'<div class="additional-item"><strong>{add.title}:</strong> {add.description or ""}</div>')
            elif add.description:
                items.append(f'<div class="additional-item">{add.description}</div>')
        if items:
            sections.append(f'<div class="section"><h2 class="section-title">Additional Information</h2>{"".join(items)}</div>')

    css = """
body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    font-size: 10pt;
    line-height: 1.45;
    color: #1e293b;
    margin: 0;
    padding: 32px 36px;
}
.header {
    text-align: center;
    margin-bottom: 6px;
}
.header h1 {
    font-size: 20pt;
    font-weight: 700;
    color: #0f172a;
    margin: 0 0 2px 0;
    letter-spacing: -0.01em;
}
.headline {
    font-size: 11pt;
    font-weight: 500;
    color: #475569;
    margin: 2px 0 0 0;
}
.contact {
    text-align: center;
    font-size: 9pt;
    color: #64748b;
    margin-bottom: 14px;
    padding-bottom: 8px;
    border-bottom: 1px solid #e2e8f0;
}
.section {
    margin-bottom: 12px;
}
.section-title {
    font-size: 10.5pt;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.06em;
    color: #0f172a;
    border-bottom: 1.5px solid #cbd5e1;
    padding-bottom: 2px;
    margin: 10px 0 6px 0;
}
.summary-text {
    font-size: 9.5pt;
    line-height: 1.5;
    color: #334155;
    margin: 3px 0 0 0;
}
.entry {
    margin-bottom: 8px;
}
.entry-header {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    font-size: 9.5pt;
    margin-bottom: 2px;
}
.entry-title {
    color: #0f172a;
}
.entry-meta {
    font-size: 9pt;
    color: #64748b;
    float: right;
}
.bullet-list {
    margin: 2px 0 4px 0;
    padding-left: 18px;
}
.bullet-list li {
    font-size: 9.5pt;
    line-height: 1.45;
    color: #334155;
    margin-bottom: 2px;
}
.skill-row {
    font-size: 9.5pt;
    line-height: 1.5;
    color: #334155;
    margin-bottom: 3px;
}
.skill-row strong {
    color: #0f172a;
}
.project-tech, .edu-detail, .lead-desc, .languages-list {
    font-size: 9pt;
    color: #475569;
    margin: 2px 0 0 0;
}
.link-item, .additional-item {
    font-size: 9pt;
    color: #334155;
    margin-bottom: 2px;
}
"""

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
{css}
</style>
</head>
<body>
{"".join(sections)}
</body>
</html>"""
    return html


class ExportService:
    """Export resume versions to PDF and DOCX."""

    def export_pdf(self, content: ResumeContent, template: str = "minimal") -> bytes:
        html = _render_html(content, template)
        try:
            import io
            import fitz  # PyMuPDF
            buf = io.BytesIO()
            # Try Story API for proper pagination
            if hasattr(fitz, 'Story'):
                writer = fitz.DocumentWriter(buf)
                story = fitz.Story(html)
                page_rect = fitz.Rect(0, 0, 595, 842)  # A4
                content_rect = page_rect + (36, 36, -36, -36)  # margins
                more = True
                while more:
                    dev = writer.begin_page(page_rect)
                    more, _ = story.place(content_rect)
                    story.draw(dev)
                    writer.end_page()
                writer.close()
                return buf.getvalue()
            else:
                # Fallback: single tall page (pre-existing behavior)
                doc = fitz.open()
                page = doc.new_page(width=595, height=842 * 3)
                page.insert_htmlbox(page.rect, html)
                pdf_bytes = doc.tobytes()
                doc.close()
                return pdf_bytes
        except Exception as exc:
            logger.error("PDF export failed: %s", exc)
            raise RuntimeError("Unable to generate PDF") from exc

    def export_docx(self, content: ResumeContent, template: str = "minimal") -> bytes:
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt  # type: ignore
        except ImportError as exc:
            raise RuntimeError("python-docx is not installed") from exc

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        profile = content.profile
        personal = profile.personal
        doc.add_heading(personal.full_name or "Candidate", level=0)
        if personal.headline:
            doc.add_paragraph(personal.headline)
        contact_parts = filter(None, [personal.email, personal.phone, personal.location, personal.website, personal.linkedin, personal.github])
        contact_line = " | ".join(contact_parts)
        if contact_line:
            doc.add_paragraph(contact_line)

        if profile.summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(profile.summary)

        if profile.experience:
            doc.add_heading("Experience", level=1)
            for exp in profile.experience:
                doc.add_heading("{role} at {company}".format(role=exp.role or "", company=exp.company or ""), level=2)
                date_range = " — ".join(filter(None, [exp.start_date, exp.end_date if not exp.current else "Present"]))
                doc.add_paragraph(date_range)
                for bullet in exp.get_all_bullet_texts():
                    doc.add_paragraph(bullet, style="List Bullet")

        if profile.internships:
            doc.add_heading("Internships", level=1)
            for exp in profile.internships:
                doc.add_heading("{role} at {company}".format(role=exp.role or "", company=exp.company or ""), level=2)
                date_range = " — ".join(filter(None, [exp.start_date, exp.end_date if not exp.current else "Present"]))
                doc.add_paragraph(date_range)
                for bullet in exp.get_all_bullet_texts():
                    doc.add_paragraph(bullet, style="List Bullet")

        if profile.projects:
            doc.add_heading("Projects", level=1)
            for proj in profile.projects:
                doc.add_heading(proj.name or "Project", level=2)
                for detail in filter(None, [proj.description, proj.problem, proj.contribution, proj.results]):
                    doc.add_paragraph(detail, style="List Bullet")

        if profile.education:
            doc.add_heading("Education", level=1)
            for edu in profile.education:
                heading = "{degree} in {field}".format(degree=edu.degree or "", field=edu.field or "")
                if edu.institution:
                    heading += " — {inst}".format(inst=edu.institution)
                doc.add_heading(heading, level=2)
                details = []
                date_range = " — ".join(filter(None, [edu.start_date, edu.end_date]))
                if date_range:
                    details.append(date_range)
                if edu.location:
                    details.append(edu.location)
                if edu.gpa:
                    details.append("GPA: {gpa}".format(gpa=edu.gpa))
                if details:
                    doc.add_paragraph(" | ".join(details))
                if edu.coursework:
                    doc.add_paragraph("Coursework: {cw}".format(cw=", ".join(edu.coursework)))
                if edu.achievements:
                    for ach in edu.achievements:
                        doc.add_paragraph(ach, style="List Bullet")

        if profile.skills and any([profile.skills.technical, profile.skills.tools, profile.skills.languages, profile.skills.databases, profile.skills.analytics, profile.skills.soft_skills]):
            doc.add_heading("Skills", level=1)
            for category, label in [
                ("technical", "Technical"), ("tools", "Tools"), ("languages", "Languages"),
                ("databases", "Databases"), ("analytics", "Analytics"), ("soft_skills", "Soft Skills"),
            ]:
                values = getattr(profile.skills, category, [])
                if values:
                    doc.add_paragraph("{label}: {values}".format(label=label, values=", ".join(values)))

        if profile.certifications:
            doc.add_heading("Certifications", level=1)
            for cert in profile.certifications:
                details = " | ".join(filter(None, [cert.name, cert.issuer, cert.date]))
                doc.add_paragraph(details)

        if profile.achievements:
            doc.add_heading("Achievements", level=1)
            for item in profile.achievements:
                doc.add_paragraph(item, style="List Bullet")

        if profile.leadership:
            doc.add_heading("Leadership", level=1)
            for lead in profile.leadership:
                header = "{role} — {org}".format(role=lead.role or "", org=lead.organization or "")
                doc.add_heading(header, level=2)
                date_range = " — ".join(filter(None, [lead.start_date, lead.end_date]))
                if date_range:
                    doc.add_paragraph(date_range)
                if lead.description:
                    doc.add_paragraph(lead.description)

        if profile.languages:
            doc.add_heading("Languages", level=1)
            items = []
            for lang in profile.languages:
                label = lang.language or ""
                if lang.proficiency:
                    label += " ({prof})".format(prof=lang.proficiency)
                items.append(label)
            doc.add_paragraph(", ".join(items))

        if profile.links:
            doc.add_heading("Links", level=1)
            for link in profile.links:
                label = link.label or link.url or ""
                doc.add_paragraph("{label}: {url}".format(label=label, url=link.url or ""))

        if profile.additional:
            doc.add_heading("Additional Information", level=1)
            for add in profile.additional:
                if add.title:
                    doc.add_paragraph("{title}: {desc}".format(title=add.title, desc=add.description or ""))
                elif add.description:
                    doc.add_paragraph(add.description)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


export_service = ExportService()
