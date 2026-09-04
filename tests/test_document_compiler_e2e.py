"""Comprehensive End-to-End Test Suite for CareerOS Resume Document Compiler.

Tests the full canonical pipeline:
REAL RESUME + REAL TARGET JD
-> Semantic Model + Geometry + Style Model
-> Senior Recruiter optimization & structured operations
-> Canonical Document Model mutation
-> Native OOXML DOCX compilation
-> PDF conversion & rendering
-> Visual verification
-> Closed-loop ATS re-analysis
-> Persistence & export artifact verification.
"""

from __future__ import annotations

import io
import fitz
import pytest
from docx import Document as DocxDocument

from app.models.resume import (
    BulletItem,
    ExperienceItem,
    PersonalInfo,
    ProjectItem,
    ResumeContent,
    ResumeProfile,
    SkillCategory,
)
from app.services.ats.ats_analyzer import ATSAnalyzer
from app.services.resume_parser.geometry import extract_document_geometry
from app.services.resumes.compiler_service import resume_compiler_service
from app.services.resumes.document_model import ResumeDocumentModel, build_document_model
from .benchmark_resume_studio import fixture_1_single_column
from app.services.resumes.docx_compiler import docx_compiler
from app.services.resumes.pdf_compiler import pdf_compiler
from app.services.resumes.style_model import DocumentStyleModel, extract_style_model
from app.services.resumes.visual_verification import VisualVerificationEngine


# Real target job description for a Principal/Senior Cloud Infrastructure Engineer
REAL_TARGET_JD = """
About the Role:
Stripe / CloudScale is seeking a Principal Infrastructure & AI Platforms Engineer to architect,
scale, and secure our global distributed cloud foundations.

Requirements:
- 8+ years designing fault-tolerant distributed systems and multi-region Kubernetes clusters.
- Deep expertise in Terraform, Kubernetes, high availability, site reliability engineering (SRE), and observability.
- Proven track record optimizing cloud infrastructure spend and reducing latency at high RPS (>200,000 req/sec).
- Experience leading cross-functional infrastructure initiatives and engineering standards.
- Strong proficiency with Go or Python, ArgoCD, Prometheus, and multi-cloud resilience.
"""


@pytest.fixture
def real_alex_morgan_resume() -> ResumeContent:
    """Real candidate resume content matching fixture_1_single_column."""
    profile = ResumeProfile(
        personal=PersonalInfo(
            full_name="Alex Morgan",
            headline="Principal Infrastructure Engineer",
            email="alex.morgan@example.com",
            phone="(555) 234-5678",
            location="San Francisco, CA",
            linkedin="linkedin.com/in/alexmorgan",
            github="github.com/alexmorgan",
        ),
        summary=(
            "Principal Infrastructure Engineer with 10+ years designing fault-tolerant cloud platforms "
            "and distributed messaging architectures."
        ),
        experience=[
            ExperienceItem(
                id="exp_001",
                role="Staff Infrastructure Architect",
                company="CloudScale Systems",
                location="San Francisco, CA",
                start_date="2020",
                end_date="Present",
                current=True,
                responsibilities=[
                    BulletItem(id="blt_001", text="Architected multi-region Kubernetes clusters supporting 250k RPS with 99.999% SLA."),
                    BulletItem(id="blt_002", text="Spearheaded database partitioning strategy saving $1.4M annually in cloud infrastructure."),
                ],
            ),
            ExperienceItem(
                id="exp_002",
                role="Senior DevOps Engineer",
                company="FinTech Innovations",
                location="San Francisco, CA",
                start_date="2017",
                end_date="2020",
                current=False,
                responsibilities=[
                    BulletItem(id="blt_003", text="Implemented zero-downtime CI/CD deployment pipelines using ArgoCD and GitHub Actions."),
                    BulletItem(id="blt_004", text="Automated multi-account AWS infrastructure provisioning using Terraform and Ansible."),
                ],
            ),
        ],
        projects=[
            ProjectItem(
                id="prj_001",
                name="MeshRouter",
                technologies=["Go", "Kubernetes", "eBPF"],
                description="High-throughput service mesh sidecar router achieving sub-millisecond p99 latency.",
            )
        ],
        skills=SkillCategory(
            technical=["Kubernetes", "AWS", "Terraform", "Docker", "Go", "Python", "ArgoCD", "Linux"],
            tools=["Prometheus", "Grafana", "Datadog", "GitHub Actions"],
        ),
    )
    return ResumeContent(profile=profile)


def test_style_model_extraction_from_real_pdf_geometry():
    """Verify DocumentStyleModel extracts authentic visual identity from real PDF geometry."""
    pdf_bytes = fixture_1_single_column()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    geom_map = extract_document_geometry(doc).to_dict()
    doc.close()

    style = extract_style_model(geom_map)
    assert style is not None
    assert style.page_width_pt == 612.0  # Letter width
    assert style.page_height_pt == 792.0  # Letter height
    assert style.margin_left_pt >= 24.0
    assert style.margin_right_pt >= 24.0
    assert style.name_size_pt >= 18.0
    assert style.body_size_pt > 0
    assert style.body_font in ("helv", "Helvetica", "Calibri", "Arial")


def test_canonical_document_model_operations(real_alex_morgan_resume):
    """Verify ResumeDocumentModel supports structured recruiter operations without layout destruction."""
    doc_model = build_document_model(real_alex_morgan_resume)
    assert doc_model.header.full_name == "Alex Morgan"
    assert len(doc_model.experience) == 2
    assert doc_model.summary is not None

    # 1. Operation: Rewrite Summary (replace_block)
    optimized_summary = (
        "Principal Infrastructure & SRE Architect with 10+ years scaling high-availability cloud platforms, "
        "multi-region Kubernetes clusters, and automated Terraform infrastructure for mission-critical enterprise systems."
    )
    ok = doc_model.apply_operation({
        "operation": "replace_block",
        "target": "summary",
        "new_content": optimized_summary,
    })
    assert ok is True
    assert doc_model.summary.text == optimized_summary

    # 2. Operation: Rewrite Bullet (rewrite_bullet)
    optimized_bullet = "Architected multi-region Kubernetes clusters supporting 250k RPS with 99.999% SLA and automated disaster recovery."
    ok = doc_model.apply_operation({
        "operation": "rewrite_bullet",
        "target": "exp_001",
        "child_id": "blt_001",
        "new_content": optimized_bullet,
    })
    assert ok is True
    assert doc_model.experience[0].bullets[0].text == optimized_bullet

    # 3. Operation: Add Bullet
    new_bullet = "Engineered enterprise observability platform utilizing Prometheus and Datadog to achieve 35% faster MTTR."
    ok = doc_model.apply_operation({
        "operation": "add_bullet",
        "target": "exp_001",
        "new_content": new_bullet,
    })
    assert ok is True
    assert len(doc_model.experience[0].bullets) == 3
    assert doc_model.experience[0].bullets[2].text == new_bullet

    # 4. Project into ResumeContent and verify lossless roundtrip
    updated_content = doc_model.to_resume_content()
    assert updated_content.profile.summary == optimized_summary
    assert updated_content.profile.experience[0].responsibilities[0].text == optimized_bullet


def test_docx_compiler_native_ooxml(real_alex_morgan_resume):
    """Verify DocxCompiler generates a native, editable Microsoft Word OOXML document."""
    doc_model = build_document_model(real_alex_morgan_resume)
    docx_bytes = docx_compiler.compile(doc_model)

    assert len(docx_bytes) > 1000
    # Must start with standard ZIP/OOXML magic bytes PK\x03\x04
    assert docx_bytes.startswith(b"PK\x03\x04")

    # Read back through python-docx to verify valid internal OOXML structure
    doc = DocxDocument(io.BytesIO(docx_bytes))
    paragraphs_text = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs_text)

    # Candidate name, sections, and bullets must exist as native Word paragraphs
    assert "Alex Morgan" in full_text
    assert "PROFESSIONAL SUMMARY" in full_text
    assert "EXPERIENCE" in full_text
    assert "CloudScale Systems" in full_text or any("CloudScale" in cell.text for t in doc.tables for row in t.rows for cell in row.cells)
    assert "Architected multi-region Kubernetes" in full_text

    # Verify native tables exist for aligned role/date headers
    assert len(doc.tables) >= 2


def test_pdf_compiler_and_visual_verification(real_alex_morgan_resume):
    """Verify PdfCompiler generates PDF and passes rigorous visual verification."""
    doc_model = build_document_model(real_alex_morgan_resume)
    docx_bytes = docx_compiler.compile(doc_model)

    pdf_bytes, ver_result = pdf_compiler.compile(doc_model, docx_bytes)
    assert len(pdf_bytes) > 1000
    assert pdf_bytes.startswith(b"%PDF-")

    # Visual verification checks
    assert ver_result.is_valid is True
    assert ver_result.page_count == 1
    assert ver_result.dimensions[0][0] > 0
    assert ver_result.dimensions[0][1] > 0

    # Extract text and verify physical presence
    fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = fitz_doc[0].get_text()
    fitz_doc.close()

    assert "Alex Morgan" in extracted_text
    assert "CloudScale Systems" in extracted_text
    assert "Kubernetes" in extracted_text


def test_real_resume_and_jd_end_to_end_closed_loop(real_alex_morgan_resume):
    """THE CRITICAL END-TO-END TEST:
    Real Resume + Real JD
    -> Parse & Geometry & Style
    -> Baseline ATS Analysis
    -> Recruiter AI structured optimization
    -> Document Model Mutation
    -> Native OOXML DOCX compilation
    -> PDF Conversion & Rendering
    -> Visual Verification
    -> Closed-loop ATS Re-analysis
    -> Verify physical content modification in generated PDF and DOCX.
    """
    # Step 1: Baseline ATS Analysis
    analyzer = ATSAnalyzer()
    initial_report = analyzer.analyze_resume(
        resume_content=real_alex_morgan_resume,
        job_description=REAL_TARGET_JD,
        job_title="Principal Infrastructure Engineer",
    )
    initial_score = initial_report.overall_score
    assert initial_score > 0

    # Step 2: Build Canonical Document Model + Style Model
    pdf_bytes = fixture_1_single_column()
    fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    orig_geom = extract_document_geometry(fitz_doc).to_dict()
    fitz_doc.close()

    doc_model = build_document_model(real_alex_morgan_resume, orig_geom)

    # Step 3: Recruiter AI Optimization (Truthful enhancement adding SRE, Terraform, Observability)
    new_summary = (
        "Principal Cloud Infrastructure & Site Reliability Engineer with 10+ years architecting "
        "high-availability distributed systems, multi-region Kubernetes clusters, and automated Terraform deployments."
    )
    new_bullet = (
        "Architected multi-region Kubernetes clusters supporting 250k RPS with 99.999% SLA, implementing "
        "automated failover, Datadog observability, and zero-downtime deployments."
    )

    # Apply structured operations
    op1_ok = doc_model.apply_operation({
        "operation": "replace_block",
        "target": "summary",
        "new_content": new_summary,
    })
    assert op1_ok is True

    op2_ok = doc_model.apply_operation({
        "operation": "rewrite_bullet",
        "target": "exp_001",
        "child_id": "blt_001",
        "new_content": new_bullet,
    })
    assert op2_ok is True

    # Step 4: Compile Native Editable DOCX
    compiled_docx_bytes = docx_compiler.compile(doc_model)
    assert len(compiled_docx_bytes) > 0

    # Step 5: Convert DOCX to PDF & visually verify
    compiled_pdf_bytes, ver_result = pdf_compiler.compile(doc_model, compiled_docx_bytes)
    assert len(compiled_pdf_bytes) > 0
    assert ver_result.is_valid is True

    # Step 6: Verify Modified Content PHYSICALLY Exists in Generated PDF
    check_pdf = fitz.open(stream=compiled_pdf_bytes, filetype="pdf")
    pdf_text = check_pdf[0].get_text()
    check_pdf.close()

    assert "Site Reliability Engineer" in pdf_text
    assert "Datadog observability" in pdf_text
    # Verify no fake placeholder exists
    assert "Your Name" not in pdf_text
    assert "Untitled Resume" not in pdf_text

    # Step 7: Verify Modified Content PHYSICALLY Exists in Generated DOCX
    check_docx = DocxDocument(io.BytesIO(compiled_docx_bytes))
    docx_text = "\n".join(p.text for p in check_docx.paragraphs)
    assert "Site Reliability Engineer" in docx_text
    assert "Datadog observability" in docx_text

    # Step 8: Closed-Loop ATS Re-analysis
    updated_resume_content = doc_model.to_resume_content()
    reanalyzed_report = analyzer.analyze_resume(
        resume_content=updated_resume_content,
        job_description=REAL_TARGET_JD,
        job_title="Principal Infrastructure Engineer",
    )
    final_score = reanalyzed_report.overall_score

    # Score must be valid and improve or stay high with genuine keyword coverage
    assert final_score >= initial_score
    initial_total_matched = len(initial_report.matched_keywords) + len(initial_report.matched_skills)
    reanalyzed_total_matched = len(reanalyzed_report.matched_keywords) + len(reanalyzed_report.matched_skills)
    assert reanalyzed_total_matched >= initial_total_matched
