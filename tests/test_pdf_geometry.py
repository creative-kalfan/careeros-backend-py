"""Comprehensive test suite for PDF Document Geometry Engine (Phase 1)."""

from __future__ import annotations

import json
import uuid
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import fitz
import pytest

from app.services.resume_parser.geometry import (
    DocumentGeometryMap,
    GeometryBlock,
    GeometryColumn,
    GeometryLine,
    GeometryPage,
    GeometrySpan,
    GeometryStyle,
    extract_document_geometry,
)
from app.services.resume_parser.layout import (
    DocumentBlock,
    DocumentLine as LayoutLine,
    DocumentSpan as LayoutSpan,
    detect_columns,
    detect_page_layout,
)
from app.services.resume_parser.models import ParseResult as ParserParseResult
from app.services.resume_parser.pdf_parser import PDFParser
from app.services.resume_parsing import ParseResult as ServiceParseResult, ResumeParsingService


def _create_single_column_pdf(path: Path) -> None:
    """Create a synthetic single-column resume with right-aligned dates."""
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    # Name and Header (spanning block)
    page.insert_text(fitz.Point(54, 50), "ALEX MORGAN", fontsize=20, fontname="helv")
    page.insert_text(
        fitz.Point(54, 75),
        "alex@example.com | (555) 019-2834 | New York, NY",
        fontsize=10,
        fontname="helv",
    )

    # Section: EXPERIENCE
    page.insert_text(fitz.Point(54, 120), "EXPERIENCE", fontsize=14, fontname="helv")

    # Job 1
    page.insert_textbox(
        fitz.Rect(54, 135, 400, 165),
        "Senior Software Engineer\nTech Innovations Corp",
        fontsize=11,
        fontname="helv",
    )
    # Right-aligned date pill
    page.insert_text(fitz.Point(450, 145), "Jan 2021 - Present", fontsize=10, fontname="helv")

    # Bullets
    page.insert_textbox(
        fitz.Rect(54, 170, 550, 230),
        "• Architected scalable distributed backend processing pipeline.\n"
        "• Mentored junior engineers and improved test coverage to 95%.\n"
        "• Optimized PostgreSQL database queries reducing p99 latency.",
        fontsize=10,
        fontname="helv",
    )

    # Job 2
    page.insert_textbox(
        fitz.Rect(54, 250, 400, 280),
        "Software Engineer\nAcme Systems Inc",
        fontsize=11,
        fontname="helv",
    )
    # Right-aligned date pill
    page.insert_text(fitz.Point(450, 260), "Jun 2018 - Dec 2020", fontsize=10, fontname="helv")

    # Bullets
    page.insert_textbox(
        fitz.Rect(54, 285, 550, 340),
        "• Developed REST and GraphQL APIs serving 2M daily active users.\n"
        "• Implemented asynchronous event streaming with Kafka and Redis.",
        fontsize=10,
        fontname="helv",
    )

    # Section: EDUCATION
    page.insert_text(fitz.Point(54, 370), "EDUCATION", fontsize=14, fontname="helv")
    page.insert_textbox(
        fitz.Rect(54, 385, 450, 420),
        "Bachelor of Science in Computer Science\nUniversity of California, Berkeley",
        fontsize=11,
        fontname="helv",
    )
    page.insert_text(fitz.Point(480, 395), "2014 - 2018", fontsize=10, fontname="helv")

    # Section: SKILLS
    page.insert_text(fitz.Point(54, 440), "SKILLS", fontsize=14, fontname="helv")
    page.insert_textbox(
        fitz.Rect(54, 455, 550, 500),
        "Languages: Python, Go, TypeScript, SQL\n"
        "Technologies: Docker, Kubernetes, AWS, FastAPI, PostgreSQL",
        fontsize=10,
        fontname="helv",
    )

    doc.save(str(path))
    doc.close()


def _create_two_column_pdf(path: Path) -> None:
    """Create a synthetic two-column resume (Left = Contact/Skills/Edu, Right = Experience)."""
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    # Full-width header at top
    page.insert_text(fitz.Point(54, 50), "SARAH CONNOR", fontsize=22, fontname="helv")

    # Left Column (x: 54 to 220)
    page.insert_text(fitz.Point(54, 100), "CONTACT", fontsize=12, fontname="helv")
    page.insert_textbox(
        fitz.Rect(54, 115, 220, 170),
        "sarah@example.com\n(555) 999-0000\nAustin, TX\nlinkedin.com/in/sconnor",
        fontsize=9,
        fontname="helv",
    )

    page.insert_text(fitz.Point(54, 200), "SKILLS", fontsize=12, fontname="helv")
    page.insert_textbox(
        fitz.Rect(54, 215, 220, 320),
        "• Python\n• React\n• Docker\n• PostgreSQL\n• FastAPI\n• Git & CI/CD",
        fontsize=9,
        fontname="helv",
    )

    page.insert_text(fitz.Point(54, 350), "EDUCATION", fontsize=12, fontname="helv")
    page.insert_textbox(
        fitz.Rect(54, 365, 220, 440),
        "BS Computer Science\nUniversity of Texas\nGraduated 2020\nGPA: 3.8",
        fontsize=9,
        fontname="helv",
    )

    # Right Column (x: 250 to 560)
    page.insert_text(fitz.Point(250, 100), "PROFESSIONAL EXPERIENCE", fontsize=14, fontname="helv")
    page.insert_textbox(
        fitz.Rect(250, 115, 560, 230),
        "Lead Fullstack Developer — Austin Cloud Systems\n"
        "2021 - Present\n"
        "• Led team of 6 engineers across 3 mission-critical cloud deployments.\n"
        "• Designed and maintained event-driven backend microservices with Kafka.\n"
        "• Increased pipeline throughput by 40% using Redis caching layers.",
        fontsize=10,
        fontname="helv",
    )

    page.insert_textbox(
        fitz.Rect(250, 250, 560, 360),
        "Software Engineer — Cyberdyne Operations\n"
        "2020 - 2021\n"
        "• Built automated telemetry ingestion pipelines using Python and async workers.\n"
        "• Implemented unit and integration test fixtures raising coverage to 92%.\n"
        "• Participated in agile ceremonies and reviewed 100+ pull requests.",
        fontsize=10,
        fontname="helv",
    )

    doc.save(str(path))
    doc.close()


class TestPDFGeometryEngine:
    """Test PDF geometry models, extraction, determinism, and layout hardening."""

    def test_geometry_single_column(self, tmp_path: Path) -> None:
        """Synthetic single-column resume with right-aligned dates.

        Verifies 1 column, bounding boxes, styles, and font preservation.
        """
        pdf_path = tmp_path / "single_column.pdf"
        _create_single_column_pdf(pdf_path)

        parser = PDFParser(debug=True)
        result = parser.parse(str(pdf_path))

        assert result.status == "completed"
        assert result.geometry is not None
        geom = result.geometry

        assert geom["page_count"] == 1
        page0 = geom["pages"][0]
        assert page0["width"] == 612.0
        assert page0["height"] == 792.0
        # Right aligned dates should NOT cause false positive 2-column layout
        assert page0["is_multi_column"] is False
        assert len(page0["columns"]) == 1

        # Check blocks exist and have valid structure
        assert len(page0["blocks"]) >= 5
        for block in page0["blocks"]:
            assert block["id"].startswith("p0_b")
            assert block["column_id"] == "p0_col0"
            assert len(block["bbox"]) == 4
            assert block["bbox"][0] < block["bbox"][2]
            assert block["bbox"][1] < block["bbox"][3]
            assert block["char_limit"] >= len(block["text"])
            assert block["style"]["font_size"] > 0
            assert len(block["lines"]) >= 1

    def test_geometry_two_column(self, tmp_path: Path) -> None:
        """Synthetic two-column resume.

        Verifies 2 detected columns with non-overlapping bounds.
        """
        pdf_path = tmp_path / "two_column.pdf"
        _create_two_column_pdf(pdf_path)

        parser = PDFParser(debug=True)
        result = parser.parse(str(pdf_path))

        assert result.status == "completed"
        assert result.geometry is not None
        geom = result.geometry

        page0 = geom["pages"][0]
        assert page0["is_multi_column"] is True
        assert len(page0["columns"]) >= 2

        col0 = page0["columns"][0]
        col1 = page0["columns"][1]

        # Non-overlapping horizontal column boundaries
        assert col0["x1"] <= col1["x0"] + 30.0
        assert col0["width"] > 80.0
        assert col1["width"] > 80.0

        # Blocks should be assigned to columns
        col_ids = {b["column_id"] for b in page0["blocks"] if b["column_id"]}
        assert col0["id"] in col_ids
        assert col1["id"] in col_ids

    def test_geometry_multi_page(self, tmp_path: Path) -> None:
        """2-page document with differing block counts and page indices."""
        pdf_path = tmp_path / "multi_page.pdf"
        doc = fitz.open()

        page1 = doc.new_page(width=612, height=792)
        page1.insert_text(fitz.Point(54, 50), "PAGE ONE HEADER", fontsize=16)
        page1.insert_textbox(
            fitz.Rect(54, 80, 500, 200),
            "Experience entry on page 1.\nMultiple lines of text.",
            fontsize=11,
        )

        page2 = doc.new_page(width=612, height=792)
        page2.insert_text(fitz.Point(54, 50), "PAGE TWO HEADER", fontsize=16)
        page2.insert_textbox(
            fitz.Rect(54, 80, 500, 150),
            "Education and Projects on page 2.",
            fontsize=11,
        )

        doc.save(str(pdf_path))
        doc.close()

        parser = PDFParser()
        result = parser.parse(str(pdf_path))

        assert result.status == "completed"
        assert result.geometry is not None
        geom = result.geometry

        assert geom["page_count"] == 2
        assert len(geom["pages"]) == 2
        assert geom["pages"][0]["page_index"] == 0
        assert geom["pages"][1]["page_index"] == 1

        # Check block ID prefixes match respective pages
        for b in geom["pages"][0]["blocks"]:
            assert b["id"].startswith("p0_b")
            assert b["page"] == 0

        for b in geom["pages"][1]["blocks"]:
            assert b["id"].startswith("p1_b")
            assert b["page"] == 1

    def test_geometry_custom_page_dimensions(self, tmp_path: Path) -> None:
        """Tests Letter (612x792) vs A4 (595.28x841.89) geometry."""
        # A4 document
        a4_path = tmp_path / "a4.pdf"
        doc = fitz.open()
        doc.new_page(width=595.3, height=841.9)
        doc[0].insert_text(fitz.Point(50, 50), "A4 Document Test", fontsize=14)
        doc.save(str(a4_path))
        doc.close()

        parser = PDFParser()
        result = parser.parse(str(a4_path))

        assert result.status == "completed"
        geom = result.geometry
        assert geom is not None
        page = geom["pages"][0]
        assert abs(page["width"] - 595.3) < 1.0
        assert abs(page["height"] - 841.9) < 1.0

    def test_geometry_missing_sections(self, tmp_path: Path) -> None:
        """Resume missing summary/projects; verifies uncertain blocks default to section=None."""
        pdf_path = tmp_path / "sparse.pdf"
        doc = fitz.open()
        page = doc.new_page(width=612, height=792)
        page.insert_text(fitz.Point(54, 50), "Jane Smith", fontsize=18)
        page.insert_text(fitz.Point(54, 70), "jane@example.com", fontsize=10)
        # Random arbitrary text that is not a known section
        page.insert_textbox(
            fitz.Rect(54, 200, 500, 300),
            "Unlabeled arbitrary paragraph of notes or hobbies.\nNot in any lexicon header.",
            fontsize=10,
        )
        doc.save(str(pdf_path))
        doc.close()

        parser = PDFParser()
        result = parser.parse(str(pdf_path))

        assert result.status == "completed"
        geom = result.geometry
        assert geom is not None
        page_blocks = geom["pages"][0]["blocks"]

        # The arbitrary block should have section=None (or ambiguous)
        arbitrary_blocks = [
            b for b in page_blocks if "Unlabeled arbitrary paragraph" in b["text"]
        ]
        assert len(arbitrary_blocks) >= 1
        assert arbitrary_blocks[0]["section"] is None

    def test_geometry_determinism(self, tmp_path: Path) -> None:
        """Parsing the same PDF twice yields identical JSON representation."""
        pdf_path = tmp_path / "deterministic.pdf"
        _create_single_column_pdf(pdf_path)

        parser = PDFParser()
        res1 = parser.parse(str(pdf_path))
        res2 = parser.parse(str(pdf_path))

        json1 = json.dumps(res1.geometry, sort_keys=True)
        json2 = json.dumps(res2.geometry, sort_keys=True)

        assert json1 == json2

    def test_geometry_validation(self, tmp_path: Path) -> None:
        """Verifies x0 < x1, y0 < y1, page bounds, non-empty text, non-zero font sizes."""
        pdf_path = tmp_path / "valid.pdf"
        _create_two_column_pdf(pdf_path)

        parser = PDFParser()
        res = parser.parse(str(pdf_path))
        geom = res.geometry
        assert geom is not None

        for page in geom["pages"]:
            pw = page["width"]
            ph = page["height"]
            for col in page["columns"]:
                assert col["x0"] < col["x1"]
                assert col["width"] > 0
                assert 0.0 <= col["x0"] <= pw
                assert 0.0 <= col["x1"] <= pw

            for b in page["blocks"]:
                assert b["bbox"][0] < b["bbox"][2]
                assert b["bbox"][1] < b["bbox"][3]
                assert 0.0 <= b["bbox"][0] <= pw
                assert 0.0 <= b["bbox"][1] <= ph
                assert len(b["text"].strip()) > 0
                assert b["style"]["font_size"] > 0
                assert b["style"]["line_height"] > 0

                for line in b["lines"]:
                    assert line["bbox"][0] <= line["bbox"][2]
                    assert line["bbox"][1] <= line["bbox"][3]
                    for s in line["spans"]:
                        assert s["size"] > 0
                        assert len(s["text"]) > 0

    @pytest.mark.asyncio
    async def test_geometry_persistence_flow(self, tmp_path: Path) -> None:
        """Verify resume_jobs and functions save geometry into meta and create_version."""
        from app.workers.functions import parse_resume_job
        from app.repositories.resume_repository import ResumeRepository

        pdf_path = tmp_path / "job_test.pdf"
        _create_single_column_pdf(pdf_path)
        pdf_bytes = pdf_path.read_bytes()

        resume_id = str(uuid.uuid4())
        user_id = str(uuid.uuid4())
        storage_path = f"{user_id}/{uuid.uuid4()}.pdf"

        mock_row = {
            "id": resume_id,
            "user_id": user_id,
            "parse_status": "pending",
            "storage_path": storage_path,
        }

        mock_repo = MagicMock(spec=ResumeRepository)
        mock_repo.get_resume.return_value = mock_row
        mock_repo.create_version.return_value = {"id": str(uuid.uuid4())}

        with patch("app.workers.functions.ResumeRepository", return_value=mock_repo):
            with patch("app.workers.functions.get_service_client") as mock_get_client:
                mock_storage = MagicMock()
                mock_storage.storage.from_.return_value.download.return_value = pdf_bytes
                mock_get_client.return_value = mock_storage

                result = await parse_resume_job(
                    ctx={"job_id": "test-job-1"},
                    resume_id=resume_id,
                    user_id=user_id,
                    storage_path=storage_path,
                )

                assert result["status"] == "completed"

                # Verify repo.update_resume was called with meta containing geometry
                update_calls = mock_repo.update_resume.call_args_list
                assert len(update_calls) >= 1
                completed_update = None
                for call in update_calls:
                    payload = call[0][2]
                    if payload.get("parse_status") == "completed":
                        completed_update = payload
                        break

                assert completed_update is not None
                assert "meta" in completed_update
                assert "geometry" in completed_update["meta"]
                assert completed_update["meta"]["geometry"] is not None
                assert completed_update["meta"]["geometry"]["page_count"] == 1

                # Verify repo.create_version was called with meta containing geometry
                create_version_calls = mock_repo.create_version.call_args_list
                assert len(create_version_calls) == 1
                version_kwargs = create_version_calls[0][1]
                assert "meta" in version_kwargs
                assert "geometry" in version_kwargs["meta"]
                assert version_kwargs["meta"]["geometry"]["page_count"] == 1

    def test_geometry_empty_or_corrupt_pdf_fails_gracefully(self, tmp_path: Path) -> None:
        """Empty or corrupt PDFs must fail gracefully with status='failed' and geometry=None."""
        parser = PDFParser()

        # 0-byte file
        empty_pdf = tmp_path / "empty.pdf"
        empty_pdf.write_bytes(b"")
        res_empty = parser.parse(str(empty_pdf))
        assert res_empty.status == "failed"
        assert res_empty.geometry is None
        assert "Cannot open empty file" in (res_empty.error or "")

        # Corrupt file
        corrupt_pdf = tmp_path / "corrupt.pdf"
        corrupt_pdf.write_bytes(b"not a valid pdf document content")
        res_corrupt = parser.parse(str(corrupt_pdf))
        assert res_corrupt.status == "failed"
        assert res_corrupt.geometry is None

        # Blank PDF without text
        blank_pdf = tmp_path / "blank.pdf"
        doc = fitz.open()
        doc.new_page(width=612, height=792)
        doc.save(str(blank_pdf))
        doc.close()
        res_blank = parser.parse(str(blank_pdf))
        assert res_blank.status == "failed"
        assert res_blank.geometry is None
        assert "No text content" in (res_blank.error or "")

    @pytest.mark.asyncio
    async def test_non_pdf_clean_handling(self, tmp_path: Path) -> None:
        """Non-PDF files (e.g. docx) return geometry=None cleanly."""
        import docx

        docx_path = tmp_path / "sample.docx"
        doc = docx.Document()
        doc.add_heading("Taylor Developer", 0)
        doc.add_paragraph("taylor@example.com | San Francisco, CA")
        doc.add_heading("Experience", 1)
        doc.add_paragraph("Software Engineer at Example Corp")
        doc.save(str(docx_path))

        service = ResumeParsingService()
        res = await service.parse_file(str(docx_path), "sample.docx")
        assert res.status == "completed"
        assert res.geometry is None

